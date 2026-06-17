"""Step implementations for hyperlink-related features."""

from __future__ import annotations

import io
from typing import Dict, Tuple

from behave import given, then, when
from behave.runner import Context

from docx import Document

from helpers import test_docx

# given ===================================================


@given("a hyperlink")
def given_a_hyperlink(context: Context):
    document = Document(test_docx("par-hyperlinks"))
    context.hyperlink = document.paragraphs[1].hyperlinks[0]


@given("a hyperlink having a URI fragment")
def given_a_hyperlink_having_a_uri_fragment(context: Context):
    document = Document(test_docx("par-hlink-frags"))
    context.hyperlink = document.paragraphs[1].hyperlinks[0]


@given("a hyperlink having address {address} and fragment {fragment}")
def given_a_hyperlink_having_address_and_fragment(
    context: Context, address: str, fragment: str
):
    paragraph_idxs: Dict[Tuple[str, str], int] = {
        ("''", "linkedBookmark"): 1,
        ("https://foo.com", "''"): 2,
        ("https://foo.com?q=bar", "''"): 3,
        ("http://foo.com/", "intro"): 4,
        ("https://foo.com?q=bar#baz", "''"): 5,
        ("court-exif.jpg", "''"): 7,
    }
    paragraph_idx = paragraph_idxs[(address, fragment)]
    document = Document(test_docx("par-hlink-frags"))
    paragraph = document.paragraphs[paragraph_idx]
    context.hyperlink = paragraph.hyperlinks[0]


@given("a hyperlink having {zero_or_more} rendered page breaks")
def given_a_hyperlink_having_rendered_page_breaks(context: Context, zero_or_more: str):
    paragraph_idx = {
        "no": 1,
        "one": 2,
    }[zero_or_more]
    document = Document(test_docx("par-hyperlinks"))
    paragraph = document.paragraphs[paragraph_idx]
    context.hyperlink = paragraph.hyperlinks[0]


@given("a hyperlink having {one_or_more} runs")
def given_a_hyperlink_having_one_or_more_runs(context: Context, one_or_more: str):
    paragraph_idx, hyperlink_idx = {
        "one": (1, 0),
        "two": (2, 1),
    }[one_or_more]
    document = Document(test_docx("par-hyperlinks"))
    paragraph = document.paragraphs[paragraph_idx]
    context.hyperlink = paragraph.hyperlinks[hyperlink_idx]


# then =====================================================


@then("hyperlink.address is the URL of the hyperlink")
def then_hyperlink_address_is_the_URL_of_the_hyperlink(context: Context):
    actual_value = context.hyperlink.address
    expected_value = "http://yahoo.com/"
    assert actual_value == expected_value, (
        f"expected: {expected_value}, got: {actual_value}"
    )


@then("hyperlink.contains_page_break is {value}")
def then_hyperlink_contains_page_break_is_value(context: Context, value: str):
    actual_value = context.hyperlink.contains_page_break
    expected_value = {"True": True, "False": False}[value]
    assert actual_value == expected_value, (
        f"expected: {expected_value}, got: {actual_value}"
    )


@then("hyperlink.fragment is the URI fragment of the hyperlink")
def then_hyperlink_fragment_is_the_URI_fragment_of_the_hyperlink(context: Context):
    actual_value = context.hyperlink.fragment
    expected_value = "linkedBookmark"
    assert actual_value == expected_value, (
        f"expected: {expected_value}, got: {actual_value}"
    )


@then("hyperlink.runs contains only Run instances")
def then_hyperlink_runs_contains_only_Run_instances(context: Context):
    actual_value = [type(item).__name__ for item in context.hyperlink.runs]
    expected_value = ["Run" for _ in context.hyperlink.runs]
    assert actual_value == expected_value, (
        f"expected: {expected_value}, got: {actual_value}"
    )


@then("hyperlink.runs has length {value}")
def then_hyperlink_runs_has_length(context: Context, value: str):
    actual_value = len(context.hyperlink.runs)
    expected_value = int(value)
    assert actual_value == expected_value, (
        f"expected: {expected_value}, got: {actual_value}"
    )


@then("hyperlink.text is the visible text of the hyperlink")
def then_hyperlink_text_is_the_visible_text_of_the_hyperlink(context: Context):
    actual_value = context.hyperlink.text
    expected_value = "awesome hyperlink"
    assert actual_value == expected_value, (
        f"expected: {expected_value}, got: {actual_value}"
    )


@then("hyperlink.url is {value}")
def then_hyperlink_url_is_value(context: Context, value: str):
    actual_value = context.hyperlink.url
    expected_value = "" if value == "''" else value
    assert actual_value == expected_value, (
        f"expected: {expected_value}, got: {actual_value}"
    )


# -- add-hyperlink steps (reuse "a document with a paragraph" given) ----------


# when ====================================================


@when('I add a hyperlink with text "{text}" and url "{url}"')
def when_i_add_a_hyperlink_with_text_and_url(context: Context, text: str, url: str):
    context.hyperlink = context.paragraph.add_hyperlink(text=text, url=url)


@when('I add an empty hyperlink with url "{url}"')
def when_i_add_an_empty_hyperlink_with_url(context: Context, url: str):
    context.hyperlink = context.paragraph.add_hyperlink(url=url)


@when('I add a run "{text}" to the hyperlink')
def when_i_add_a_run_to_the_hyperlink(context: Context, text: str):
    context.hyperlink.add_run(text)


# then ====================================================


@then("the saved document has {count:d} hyperlink")
@then("the saved document has {count:d} hyperlinks")
def then_the_saved_document_has_n_hyperlinks(context: Context, count: int):
    buffer = io.BytesIO()
    context.document.save(buffer)
    buffer.seek(0)
    context.reopened = Document(buffer)
    context.hyperlinks = [
        hyperlink
        for paragraph in context.reopened.paragraphs
        for hyperlink in paragraph.hyperlinks
    ]
    assert len(context.hyperlinks) == count, "got %d hyperlinks" % len(
        context.hyperlinks
    )


@then('hyperlink {idx:d} has text "{text}"')
def then_hyperlink_n_has_text(context: Context, idx: int, text: str):
    actual = context.hyperlinks[idx - 1].text
    assert actual == text, "hyperlink text was %r" % actual


@then('hyperlink {idx:d} has url "{url}"')
def then_hyperlink_n_has_url(context: Context, idx: int, url: str):
    actual = context.hyperlinks[idx - 1].url
    assert actual == url, "hyperlink url was %r" % actual


@then('the saved document defines the "{style_id}" character style')
def then_the_saved_document_defines_the_character_style(
    context: Context, style_id: str
):
    styles = context.reopened.part._styles_part.element
    assert styles.get_by_id(style_id) is not None, "%r style not defined" % style_id
