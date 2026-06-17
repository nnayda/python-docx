"""Step implementations for footnote-related features."""

from __future__ import annotations

import io

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# given ===================================================


@given("a document with a paragraph")
def given_a_document_with_a_paragraph(context: Context):
    context.document = Document()
    context.paragraph = context.document.add_paragraph("body text")


@given("a document with a paragraph containing a run")
def given_a_document_with_a_paragraph_containing_a_run(context: Context):
    context.document = Document()
    context.paragraph = context.document.add_paragraph()
    context.run = context.paragraph.add_run("anchor")


# when ====================================================


@when('I add a footnote with text "{text}"')
def when_i_add_a_footnote_with_text(context: Context, text: str):
    context.paragraph.add_footnote(text)


@when('I add a footnote to the run with text "{text}"')
def when_i_add_a_footnote_to_the_run_with_text(context: Context, text: str):
    context.run.add_footnote(text)


# then ====================================================


@then("the saved document has {count:d} footnote")
@then("the saved document has {count:d} footnotes")
def then_the_saved_document_has_n_footnotes(context: Context, count: int):
    buffer = io.BytesIO()
    context.document.save(buffer)
    buffer.seek(0)
    context.reopened = Document(buffer)
    footnotes_part = context.reopened.part.part_related_by(RT.FOOTNOTES)
    # -- user footnotes exclude the two separator entries (ids -1 and 0) --
    user_footnotes = [ftn for ftn in footnotes_part.element.footnote_lst if ftn.id > 0]
    context.footnotes = user_footnotes
    assert len(user_footnotes) == count, "got %d footnotes" % len(user_footnotes)


@then('footnote {idx:d} has text "{text}"')
def then_footnote_n_has_text(context: Context, idx: int, text: str):
    ftn = next(f for f in context.footnotes if f.id == idx)
    actual = "".join(p.text for p in ftn.p_lst)
    assert text in actual, "footnote text was %r" % actual
