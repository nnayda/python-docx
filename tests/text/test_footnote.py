# pyright: reportPrivateUsage=false

"""Test suite for the docx.text.footnote module."""

from __future__ import annotations

from docx import Document
from docx.text.footnote import Footnote
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def _new_footnote(text: str | None = None) -> Footnote:
    """Return a Footnote freshly added to a blank document's footnotes part."""
    return Document().part.add_footnote(text)


class DescribeFootnote:
    """Unit-test suite for the Footnote proxy."""

    def it_knows_its_id(self) -> None:
        footnote = _new_footnote("note")
        assert footnote.id == 1

    def it_knows_its_text(self) -> None:
        footnote = _new_footnote("the reference text")
        assert footnote.text == "the reference text"

    def it_starts_empty_when_no_text_given(self) -> None:
        footnote = _new_footnote()
        assert footnote.text == ""

    def it_can_replace_its_text(self) -> None:
        footnote = _new_footnote("old")
        footnote.text = "new"
        assert footnote.text == "new"

    def it_can_add_a_paragraph_to_its_body(self) -> None:
        footnote = _new_footnote()
        paragraph = footnote.add_paragraph("second line")
        assert isinstance(paragraph, Paragraph)
        assert paragraph.style is not None
        assert paragraph.style.style_id == "FootnoteText"
        assert footnote.text.endswith("second line")

    def it_can_add_a_run_to_its_last_paragraph(self) -> None:
        footnote = _new_footnote("start")
        run = footnote.add_run(" more")
        assert isinstance(run, Run)
        assert footnote.text == "start more"

    def it_begins_with_a_footnote_ref_mark(self) -> None:
        footnote = _new_footnote("x")
        assert footnote._element.xpath("./w:p[1]/w:r[1]/w:footnoteRef")
