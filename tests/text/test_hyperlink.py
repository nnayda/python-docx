"""Test suite for the docx.text.hyperlink module."""

from typing import cast

import pytest

from docx import Document
from docx import types as t
from docx.opc.rel import _Relationship  # pyright: ignore[reportPrivateUsage]
from docx.oxml.text.hyperlink import CT_Hyperlink
from docx.parts.story import StoryPart
from docx.text.hyperlink import Hyperlink
from docx.text.run import Run

from ..unitutil.cxml import element
from ..unitutil.mock import FixtureRequest, Mock, instance_mock


class DescribeHyperlink:
    """Unit-test suite for the docx.text.hyperlink.Hyperlink object."""

    @pytest.mark.parametrize(
        ("hlink_cxml", "expected_value"),
        [
            ('w:hyperlink{r:id=rId6}/w:r/w:t"post"', "https://google.com/"),
            ("w:hyperlink{w:anchor=_Toc147925734}", ""),
            ("w:hyperlink", ""),
        ],
    )
    def it_knows_the_hyperlink_address(
        self, hlink_cxml: str, expected_value: str, fake_parent: t.ProvidesStoryPart
    ):
        hlink = cast(CT_Hyperlink, element(hlink_cxml))
        hyperlink = Hyperlink(hlink, fake_parent)

        assert hyperlink.address == expected_value

    @pytest.mark.parametrize(
        ("hlink_cxml", "expected_value"),
        [
            ("w:hyperlink", False),
            ("w:hyperlink/w:r", False),
            ('w:hyperlink/w:r/(w:t"abc",w:lastRenderedPageBreak,w:t"def")', True),
            ('w:hyperlink/w:r/(w:lastRenderedPageBreak,w:t"abc",w:t"def")', True),
            ('w:hyperlink/w:r/(w:t"abc",w:t"def",w:lastRenderedPageBreak)', True),
        ],
    )
    def it_knows_whether_it_contains_a_page_break(
        self, hlink_cxml: str, expected_value: bool, fake_parent: t.ProvidesStoryPart
    ):
        hlink = cast(CT_Hyperlink, element(hlink_cxml))
        hyperlink = Hyperlink(hlink, fake_parent)

        assert hyperlink.contains_page_break is expected_value

    @pytest.mark.parametrize(
        ("hlink_cxml", "expected_value"),
        [
            ("w:hyperlink{r:id=rId6}", ""),
            ("w:hyperlink{w:anchor=intro}", "intro"),
        ],
    )
    def it_knows_the_link_fragment_when_there_is_one(
        self, hlink_cxml: str, expected_value: str, fake_parent: t.ProvidesStoryPart
    ):
        hlink = cast(CT_Hyperlink, element(hlink_cxml))
        hyperlink = Hyperlink(hlink, fake_parent)

        assert hyperlink.fragment == expected_value

    @pytest.mark.parametrize(
        ("hlink_cxml", "count"),
        [
            ("w:hyperlink", 0),
            ("w:hyperlink/w:r", 1),
            ("w:hyperlink/(w:r,w:r)", 2),
            ("w:hyperlink/(w:r,w:lastRenderedPageBreak)", 1),
            ("w:hyperlink/(w:lastRenderedPageBreak,w:r)", 1),
            ("w:hyperlink/(w:r,w:lastRenderedPageBreak,w:r)", 2),
        ],
    )
    def it_provides_access_to_the_runs_it_contains(
        self, hlink_cxml: str, count: int, fake_parent: t.ProvidesStoryPart
    ):
        hlink = cast(CT_Hyperlink, element(hlink_cxml))
        hyperlink = Hyperlink(hlink, fake_parent)

        runs = hyperlink.runs

        actual = [type(item).__name__ for item in runs]
        expected = ["Run" for _ in range(count)]
        assert actual == expected

    @pytest.mark.parametrize(
        ("hlink_cxml", "expected_text"),
        [
            ("w:hyperlink", ""),
            ("w:hyperlink/w:r", ""),
            ('w:hyperlink/w:r/w:t"foobar"', "foobar"),
            ('w:hyperlink/w:r/(w:t"foo",w:lastRenderedPageBreak,w:t"bar")', "foobar"),
            ('w:hyperlink/w:r/(w:t"abc",w:tab,w:t"def",w:noBreakHyphen)', "abc\tdef-"),
        ],
    )
    def it_knows_the_visible_text_of_the_link(
        self, hlink_cxml: str, expected_text: str, fake_parent: t.ProvidesStoryPart
    ):
        hlink = cast(CT_Hyperlink, element(hlink_cxml))
        hyperlink = Hyperlink(hlink, fake_parent)

        text = hyperlink.text

        assert text == expected_text

    @pytest.mark.parametrize(
        ("hlink_cxml", "expected_value"),
        [
            ("w:hyperlink", ""),
            ("w:hyperlink{w:anchor=_Toc147925734}", ""),
            ('w:hyperlink{r:id=rId6}/w:r/w:t"post"', "https://google.com/"),
            (
                'w:hyperlink{r:id=rId6,w:anchor=foo}/w:r/w:t"post"',
                "https://google.com/#foo",
            ),
        ],
    )
    def it_knows_the_full_url_for_web_addresses(
        self, hlink_cxml: str, expected_value: str, fake_parent: t.ProvidesStoryPart
    ):
        hlink = cast(CT_Hyperlink, element(hlink_cxml))
        hyperlink = Hyperlink(hlink, fake_parent)

        assert hyperlink.url == expected_value

    def it_can_add_a_run_carrying_the_built_in_Hyperlink_style(self):
        hyperlink = self._empty_hyperlink()

        run = hyperlink.add_run("click here")

        assert isinstance(run, Run)
        assert run.text == "click here"
        assert run._r.style == "Hyperlink"  # pyright: ignore[reportPrivateUsage]
        assert [r.text for r in hyperlink.runs] == ["click here"]

    def it_can_add_an_empty_run(self):
        hyperlink = self._empty_hyperlink()

        run = hyperlink.add_run()

        assert run.text == ""
        assert run._r.style == "Hyperlink"  # pyright: ignore[reportPrivateUsage]

    def it_can_add_a_run_with_an_overriding_style(self):
        hyperlink = self._empty_hyperlink()

        run = hyperlink.add_run("bold link", style="Strong")

        assert run._r.style == "Strong"  # pyright: ignore[reportPrivateUsage]

    @staticmethod
    def _empty_hyperlink() -> Hyperlink:
        """Return a |Hyperlink| whose part is a real (blank) document part."""
        paragraph = Document().add_paragraph()
        hlink = cast(CT_Hyperlink, element("w:hyperlink{r:id=rId6}"))
        return Hyperlink(hlink, paragraph)

    # -- fixtures --------------------------------------------------------------------

    @pytest.fixture
    def fake_parent(self, story_part: Mock, rel: Mock) -> t.ProvidesStoryPart:
        class StoryChild:
            @property
            def part(self) -> StoryPart:
                return story_part

        return StoryChild()

    @pytest.fixture
    def rel(self, request: FixtureRequest):
        return instance_mock(request, _Relationship, target_ref="https://google.com/")

    @pytest.fixture
    def story_part(self, request: FixtureRequest, rel: Mock):
        return instance_mock(request, StoryPart, rels={"rId6": rel})
