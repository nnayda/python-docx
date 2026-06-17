"""Round-trip integration test for footnote creation."""

import io
from typing import cast

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.footnotes import FootnotesPart


class DescribeFootnoteRoundTrip:
    """End-to-end: create footnotes, save, reopen, verify."""

    def it_persists_footnotes_through_save_and_reopen(self):
        document = Document()
        p = document.add_paragraph("As Smith showed")
        r = p.add_run(", refined later")
        r.add_footnote("Smith, J. (2024). Title. Publisher.")
        p.add_footnote("Jones, A. (2023). Another work.")

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        reopened = Document(buffer)

        footnotes_part = cast(
            FootnotesPart, reopened.part.part_related_by(RT.FOOTNOTES)
        )
        footnotes = footnotes_part.element
        # -- two separators (-1, 0) plus two user footnotes (1, 2) --
        ids = [ftn.id for ftn in footnotes.footnote_lst]
        assert ids == [-1, 0, 1, 2]
        texts = ["".join(p.text for p in ftn.p_lst) for ftn in footnotes.footnote_lst]
        assert "Smith, J. (2024). Title. Publisher." in texts[2]
        assert "Jones, A. (2023). Another work." in texts[3]

    def it_writes_reference_markers_into_the_body(self):
        document = Document()
        p = document.add_paragraph("body")
        p.add_footnote("note")

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        reopened = Document(buffer)

        body_refs = reopened.part.element.xpath(".//w:footnoteReference/@w:id")
        assert "1" in body_refs
